"""导出服务：export_docx.py 的数据库适配版。

章节来源 = chapter_versions 最新版本行；大纲一级章标题 = outline_snapshots；
渲染逻辑（样式/封面/目录/表格/[待补]高亮/标题层级映射）直接复用 scripts/export_docx.py。
"""
import sys
from datetime import datetime, timezone
from pathlib import Path

from sqlalchemy.orm import Session

sys.path.insert(0, str(Path(__file__).resolve().parents[2] / "scripts"))

from docx import Document  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.shared import Pt  # noqa: E402

from app.core.database import SessionLocal  # noqa: E402
from app.core.storage import storage  # noqa: E402
from app.models.chapter import Chapter, ChapterVersion  # noqa: E402
from app.models.outline import OutlineSnapshot  # noqa: E402
from app.models.project import Project  # noqa: E402

from scripts.export_docx import (  # noqa: E402
    add_page_number_footer,
    add_toc,
    chapter_sort_key,
    render_markdown,
    set_cjk_font,
    set_update_fields_on_open,
    setup_styles,
)


def _chapter_sort(key: str):
    return chapter_sort_key(f"{key}-x.md")


def build_export_doc(project_id: int, db: Session):
    """组装 docx 文档对象（导出与预览共用）。返回 (doc, meta) 或 (None, error)。"""
    project = db.get(Project, project_id)
    if project is None:
        return None, {"error": "项目不存在"}

    chapters = (
        db.query(Chapter)
        .filter(Chapter.project_id == project_id)
        .all()
    )
    contents: dict[str, tuple[str, str]] = {}  # key -> (title, content)
    total_words = 0
    pending = 0
    for ch in chapters:
        v = (
            db.query(ChapterVersion)
            .filter(ChapterVersion.chapter_id == ch.id)
            .order_by(ChapterVersion.version_no.desc())
            .first()
        )
        if v and v.content_md.strip():
            contents[ch.chapter_key] = (ch.title, v.content_md)
            total_words += v.word_count
            pending += v.content_md.count("[待补")
    if not contents:
        return None, {"error": "尚无章节正文可导出"}

    snap = (
        db.query(OutlineSnapshot)
        .filter(OutlineSnapshot.project_id == project_id, OutlineSnapshot.version == project.outline_version)
        .first()
    )
    parent_titles = {}
    if snap:
        for n in snap.tree.get("nodes", []):
            parent_titles[str(n["id"])] = n.get("title", "")

    doc = Document()
    setup_styles(doc)
    add_page_number_footer(doc)
    set_update_fields_on_open(doc)

    # 封面
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("投 标 文 件"); r.bold = True; r.font.size = Pt(36); set_cjk_font(r, "黑体")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("（技术文件）"); r.bold = True; r.font.size = Pt(22); set_cjk_font(r, "黑体")
    for _ in range(3):
        doc.add_paragraph()
    for label, val in [("项目名称", project.name), ("招标编号", project.tender_no), ("投标人", "（盖章）"), ("日期", "")]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{label}：{val}"); r.font.size = Pt(14); set_cjk_font(r)
    doc.add_page_break()

    # 目录
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("目  录"); r.bold = True; r.font.size = Pt(16); set_cjk_font(r, "黑体")
    add_toc(doc)
    doc.add_page_break()

    # 章节（按大纲 id 数值排序；子节父章无正文时补发父章标题）
    emitted_parents: set[str] = set()
    keys = sorted(contents.keys(), key=_chapter_sort)
    top_keys = {k for k in keys if "." not in k}
    for key in keys:
        title, text = contents[key]
        if "." in key:
            parent = key.split(".")[0]
            if parent not in emitted_parents and parent not in top_keys:
                ptitle = parent_titles.get(parent, "")
                doc.add_heading(f"{parent} {ptitle}".strip(), level=1)
                emitted_parents.add(parent)
        render_markdown(doc, text, base_id=key, ws=None)
        doc.add_page_break()

    meta = {
        "chapters": len(contents),
        "total_words": total_words,
        "pending_gaps": pending,
        "project_name": project.name,
        "tender_no": project.tender_no,
    }
    return doc, meta


def run_export_preview(project_id: int) -> dict:
    """导出前预览：结构摘要（章节清单/字数/待补/样式说明），不生成文件。"""
    db: Session = SessionLocal()
    try:
        project = db.get(Project, project_id)
        if project is None or project.state not in ("draft_done", "checking", "exported"):
            return {"error": f"状态 {project and project.state} 不允许导出预览"}
        chapters = (
            db.query(Chapter)
            .filter(Chapter.project_id == project_id)
            .order_by(Chapter.sort_order)
            .all()
        )
        items = []
        total_words = 0
        pending = 0
        for ch in chapters:
            v = (
                db.query(ChapterVersion)
                .filter(ChapterVersion.chapter_id == ch.id)
                .order_by(ChapterVersion.version_no.desc())
                .first()
            )
            if v and v.content_md.strip():
                total_words += v.word_count
                p_cnt = v.content_md.count("[待补")
                pending += p_cnt
                items.append({
                    "key": ch.chapter_key, "title": ch.title,
                    "words": v.word_count, "pending": p_cnt,
                })
        if not items:
            return {"error": "尚无章节正文可导出"}
        return {
            "project_name": project.name,
            "tender_no": project.tender_no,
            "chapters": items,
            "total_words": total_words,
            "pending_gaps": pending,
            "style_notes": [
                "封面：投标文件（技术文件）+ 项目名称/招标编号",
                "目录：自动目录（打开文档自动更新页码），宋体小四",
                "一级章标题：宋体四号加粗；二级及以下：黑体",
                "正文：宋体小四，1.3 倍行距，首行缩进 2 字符",
                "表格：Table Grid，表头加粗，表格文字五号",
                "[待补]：黄色高亮",
            ],
        }
    finally:
        db.close()


def run_export(project_id: int) -> dict:
    """合成 docx 终稿落存储。状态边：draft_done/checking → exported。"""
    db: Session = SessionLocal()
    try:
        project = db.get(Project, project_id)
        if project is None or project.state not in ("draft_done", "checking", "exported"):
            return {"error": f"状态 {project and project.state} 不允许导出"}
        doc, meta = build_export_doc(project_id, db)
        if doc is None:
            return meta
        export_rel = f"projects/{project_id}/export/技术文件.docx"
        tmp = storage.abspath(export_rel)
        tmp.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(tmp))

        project.state = "exported"
        db.commit()
        return {
            "export_path": export_rel,
            "chapters": meta["chapters"],
            "total_words": meta["total_words"],
            "pending_gaps": meta["pending_gaps"],
            "exported_at": datetime.now(timezone.utc).isoformat(),
        }
    finally:
        db.close()
