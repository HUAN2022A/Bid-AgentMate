#!/usr/bin/env python3
"""把 chapters/*.md 合成导出为技术文件 docx 终稿。

按 bid-outline.yaml 的一级章顺序排列章节，解析 markdown（标题/表格/加粗/列表），
剥离 covers 注释，[待补] 高亮，生成带封面+自动目录的 docx。
"""
import argparse
import re
import sys
from pathlib import Path

try:
    import yaml
except ImportError:
    sys.exit("缺少 PyYAML，请先: pip install pyyaml")

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor


# ---------- 样式 ----------

def set_cjk_font(run, font_cn="宋体", font_en="Times New Roman"):
    run.font.name = font_en
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), font_cn)


def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)  # 小四
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.3

    # 标题样式
    for lvl, (size, cn) in {1: (16, "黑体"), 2: (14, "黑体"), 3: (13, "黑体"), 4: (12, "黑体")}.items():
        st = doc.styles[f"Heading {lvl}"]
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.font.name = "Times New Roman"
        st.element.rPr.rFonts.set(qn("w:eastAsia"), cn)


def add_toc(doc):
    """插入自动目录（复杂域字符写法，兼容性更好）。"""
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar"); fld_sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "（目录：打开文档时自动更新，或全选按 F9 生成）"
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    for el in (fld_begin, instr, fld_sep, t, fld_end):
        run._r.append(el)


def set_update_fields_on_open(doc):
    """设置打开文档时自动更新域（目录页码自动生成，无需手动 F9）。"""
    settings = doc.settings.element
    uf = OxmlElement("w:updateFields")
    uf.set(qn("w:val"), "true")
    settings.append(uf)


def add_page_number_footer(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)


# ---------- markdown 解析 ----------

def add_runs_with_bold(p, text, highlight_pending=True):
    """处理 **加粗** 与 [待补] 高亮。"""
    # 先按 [待补...] 切分
    parts = re.split(r"(\[待补[^]]*\])", text)
    for part in parts:
        if not part:
            continue
        is_pending = part.startswith("[待补")
        # 再处理加粗
        for seg in re.split(r"(\*\*[^*]+\*\*)", part):
            if not seg:
                continue
            bold = seg.startswith("**") and seg.endswith("**")
            clean = seg[2:-2] if bold else seg
            run = p.add_run(clean)
            run.bold = bold
            set_cjk_font(run)
            if is_pending:
                # 黄色高亮
                hl = OxmlElement("w:highlight")
                hl.set(qn("w:val"), "yellow")
                run._element.get_or_add_rPr().append(hl)


def add_table(doc, rows):
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            c = t.cell(i, j)
            c.text = ""
            p = c.paragraphs[0]
            add_runs_with_bold(p, cell.strip())
            if i == 0:
                for r in p.runs:
                    r.bold = True
    return t


def render_markdown(doc, md_text, base_id="", ws=None):
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        # 跳过 covers 注释
        if line.strip().startswith("<!--"):
            while i < len(lines) and "-->" not in lines[i]:
                i += 1
            i += 1
            continue
        if not line.strip():
            i += 1
            continue
        # 图片：![图注](path)
        img = re.match(r"^!\[(.*?)\]\((.*?)\)", line.strip())
        if img:
            caption, imgpath = img.group(1), img.group(2)
            full = (ws / imgpath) if ws else Path(imgpath)
            if full.exists():
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(str(full), width=Pt(420))
                cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cr = cp.add_run(caption); cr.font.size = Pt(10.5); cr.bold = True
                set_cjk_font(cr)
            else:
                p = doc.add_paragraph(); add_runs_with_bold(p, f"[图片缺失：{imgpath}]")
            i += 1
            continue
        # 标题（层级 = 章节深度 + md 层级）
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            level = heading_level(base_id, len(m.group(1)))
            doc.add_heading(m.group(2).strip(), level=level)
            i += 1
            continue
        # 表格
        if line.strip().startswith("|"):
            tbl = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row = [c for c in lines[i].strip().strip("|").split("|")]
                # 跳过分隔行 |---|
                if not all(set(c.strip()) <= set(":- ") for c in row):
                    tbl.append(row)
                i += 1
            if tbl:
                add_table(doc, tbl)
            continue
        # 列表
        if re.match(r"^\s*[-*]\s+", line):
            text = re.sub(r"^\s*[-*]\s+", "", line)
            p = doc.add_paragraph(style="List Bullet")
            add_runs_with_bold(p, text)
            i += 1
            continue
        if re.match(r"^\s*\d+\.\s+", line):
            text = re.sub(r"^\s*\d+\.\s+", "", line)
            p = doc.add_paragraph(style="List Number")
            add_runs_with_bold(p, text)
            i += 1
            continue
        # 普通段落（首行缩进）
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(24)  # 2字符
        add_runs_with_bold(p, line)
        i += 1


# ---------- 主流程 ----------

def chapter_sort_key(fname):
    """按 outline 章节 id 排序：'1' < '2.1' < '2.10' < '10'（数值元组，非字典序）。"""
    m = re.match(r"^([\d.]+)", fname)
    if not m:
        return (999,)
    return tuple(int(x) for x in m.group(1).rstrip(".").split("."))


def heading_level(chapter_id, md_level):
    """md 的 # 层级 + 章节 id 深度 → Word Heading 级别。

    一级章（id 无点，如 '1'、'10'）的 # → H1；
    二级节（id 如 '2.1'）的 # → H2，其 ## → H3；
    依此类推，封顶 H4。
    """
    depth = chapter_id.count(".")  # '1'->0, '2.1'->1, '2.1.1'->2
    return min(depth + md_level, 4)


def chapter_id_of(fname):
    m = re.match(r"^([\d.]+)", fname)
    return m.group(1).rstrip(".") if m else ""


def main():
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")

    ap = argparse.ArgumentParser()
    ap.add_argument("-d", "--dir", required=True, help="项目工作区目录")
    ap.add_argument("-o", "--output", default=None, help="输出 docx 路径")
    args = ap.parse_args()
    ws = Path(args.dir)

    tender = yaml.safe_load((ws / "tender-analysis.yaml").read_text(encoding="utf-8")) or {}
    project = tender.get("project", "技术文件")
    tender_no = tender.get("tender_no", "")

    # 从 outline 取一级章标题（用于给无独立文件的父章补发标题）
    outline = yaml.safe_load((ws / "bid-outline.yaml").read_text(encoding="utf-8")) or {}
    parent_titles = {str(ch.get("id")): ch.get("title", "") for ch in outline.get("chapters", [])}

    out = Path(args.output) if args.output else ws / "技术文件.docx"

    doc = Document()
    setup_styles(doc)
    add_page_number_footer(doc)
    set_update_fields_on_open(doc)

    # 封面
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("投 标 文 件"); r.bold = True; r.font.size = Pt(36); set_cjk_font(r, "黑体")
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("（技术文件）"); r.bold = True; r.font.size = Pt(22); set_cjk_font(r, "黑体")
    for _ in range(3):
        doc.add_paragraph()
    for label, val in [("项目名称", project), ("招标编号", tender_no), ("投标人", "（盖章）"), ("日期", "")]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{label}：{val}"); r.font.size = Pt(14); set_cjk_font(r)
    doc.add_page_break()

    # 目录
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("目  录"); r.bold = True; r.font.size = Pt(16); set_cjk_font(r, "黑体")
    add_toc(doc)
    doc.add_page_break()

    # 章节
    chdir = ws / "chapters"
    files = sorted(chdir.glob("*.md"), key=lambda f: chapter_sort_key(f.name))
    total_words = 0
    pending = 0
    emitted_parents = set()  # 已补发的父章标题 id
    for f in files:
        text = f.read_text(encoding="utf-8")
        pending += text.count("[待补")
        body = re.sub(r"<!--.*?-->", "", text, flags=re.S)
        total_words += len(body.replace("\n", "").replace(" ", ""))
        cid = chapter_id_of(f.name)
        # 若是子节（如 2.1）且父章（2）无独立文件，先补发父章标题
        if "." in cid:
            parent = cid.split(".")[0]
            parent_has_file = any(x.name.startswith(parent + "-") for x in files)
            if parent not in emitted_parents and not parent_has_file:
                ptitle = parent_titles.get(parent, "")
                doc.add_heading(f"{parent} {ptitle}".strip(), level=1)
                emitted_parents.add(parent)
        render_markdown(doc, text, base_id=cid, ws=ws)
        doc.add_page_break()

    doc.save(str(out))
    print(f"已生成: {out}")
    print(f"章节文件数: {len(files)}，总字数约: {total_words}，[待补]高亮: {pending} 处")
    print("提示：目录已设为打开文档时自动更新页码（若未更新，全选按 F9）。")


if __name__ == "__main__":
    main()
